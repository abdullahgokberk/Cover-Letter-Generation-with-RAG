from docx import Document
from langchain.schema import Document as LangchainDocument
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings
from dotenv import load_dotenv

load_dotenv()

def read_docx(file_path):
    doc = Document(file_path)
    full_text = ""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:  # boş paragrafları atla
            #print(text)
            full_text += text + "\n"
    return full_text

docx_path = r"C:\Users\abdul\Desktop\belgelerim\CV\CvEng.docx"
print("--Reading Docx--")
#print(f"--Reading Docx--: {docx_path}")
text = read_docx(docx_path)
langchain_doc = [LangchainDocument(page_content=text)]
vectorstore = Chroma.from_documents(
    documents=langchain_doc,
    embedding=OpenAIEmbeddings(),
    collection_name="rag_cv_chroma",
    persist_directory="./.rag_cv"
)
retriever = Chroma(
    collection_name="rag_cv_chroma",
    persist_directory="./.rag_cv",
    embedding_function=OpenAIEmbeddings(),
).as_retriever()



